"""
docx_exporter.py — Exportador de relatórios para DOCX via python-docx.

Dependência opcional: pip install python-docx
Se python-docx não estiver instalado, o exporter emite um aviso e retorna None.
"""
import logging
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger("docx_exporter")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning(
        "DOCXExporter: python-docx não está instalado. "
        "Execute `pip install python-docx` para habilitar exportação DOCX."
    )


class DOCXExporter:
    """
    Converte um relatório Markdown (string) em um arquivo DOCX estruturado usando python-docx.
    Opera em modo degradado (retorna None com warning) se python-docx não estiver disponível.
    """

    def __init__(self):
        self.available = _DOCX_AVAILABLE

    def export(self, markdown_content: str, filepath: str) -> Optional[str]:
        """
        Gera o DOCX a partir do conteúdo Markdown e o salva em `filepath`.

        Returns:
            Caminho do arquivo gerado, ou None se python-docx não disponível.
        """
        if not self.available:
            logger.warning("DOCXExporter: exportação ignorada — python-docx não disponível.")
            return None

        docx_path = str(filepath).replace(".md", ".docx")
        Path(docx_path).parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Configurações de página
        section = doc.sections[0]
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(90)

        for line in markdown_content.splitlines():
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph("")
                continue

            if stripped.startswith("# "):
                doc.add_heading(self._strip_md(stripped[2:]), level=1)
            elif stripped.startswith("## "):
                doc.add_heading(self._strip_md(stripped[3:]), level=2)
            elif stripped.startswith("### "):
                doc.add_heading(self._strip_md(stripped[4:]), level=3)
            elif stripped.startswith("---"):
                # Linha horizontal — paragraph com borda inferior
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr = pPr.get_or_add_pBdr()
                bottom = pBdr.get_or_add_bottom()
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "AAAAAA")
            elif stripped.startswith("> "):
                # Blockquote como parágrafo indentado
                p = doc.add_paragraph(self._strip_md(stripped[2:]))
                p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
                p.paragraph_format.left_indent = Pt(36)
            elif stripped.startswith("- ") or re.match(r"^\d+\.", stripped):
                # Lista não-numerada ou numerada
                text = stripped.lstrip("- ").lstrip()
                text = re.sub(r"^\d+\.\s+", "", text)
                doc.add_paragraph(self._strip_md(text), style="List Bullet")
            else:
                p = doc.add_paragraph(self._strip_md(stripped))

        try:
            doc.save(docx_path)
            logger.info(f"DOCXExporter: DOCX gerado em {docx_path}")
            return docx_path
        except Exception as e:
            logger.error(f"DOCXExporter: falha ao gerar DOCX: {e}")
            return None

    def _strip_md(self, text: str) -> str:
        """
        Remove marcação Markdown básica para texto limpo no DOCX.
        """
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1 (\2)", text)
        return text
